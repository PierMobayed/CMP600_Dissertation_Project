"""
Rebuild Requirements + Evaluation Plan; convert .md/.txt in Documentation/ to .docx.
Run: python Documentation/build_all_docs.py
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt

DOC_DIR = Path(__file__).parent


def _style(doc: Document) -> None:
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(11)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for j, h in enumerate(headers):
        t.rows[0].cells[j].text = h
    for i, row in enumerate(rows, 1):
        for j, cell in enumerate(row):
            t.rows[i].cells[j].text = cell
    doc.add_paragraph()


def build_requirements() -> None:
    doc = Document()
    _style(doc)
    doc.add_heading("Requirements Document (Version 1)", 0)
    doc.add_paragraph(
        "Door2Door logistics prototype (CMP600). Scope is an MVP with simulated London data only. "
        "This register matches the submission build (May 2026)."
    )

    doc.add_heading("1. Stakeholders", level=1)
    add_table(
        doc,
        ("Role", "Need", "Prototype support"),
        [
            ("Client", "Book and track parcels", "Booking wizard, tracking map"),
            ("Driver", "Run ordered stops", "Queue, map, external navigation links"),
            ("Office dispatcher", "See delays and assign work", "KPI row, map, filters, dispatch list"),
            ("Assessor", "Verify FR/NFR evidence", "Tests, screenshots, OpenAPI"),
        ],
    )

    doc.add_heading("2. Functional requirements (implemented)", level=1)
    add_table(
        doc,
        ("ID", "Requirement", "Evidence"),
        [
            ("FR-C1", "Client creates door-to-door shipment with pickup and delivery", "client_app booking; POST shipments"),
            ("FR-C2", "Client tracks shipment on map", "GET tracking; client-tracking.png"),
            ("FR-D1", "Driver sees ordered stops", "delivery-queue; driver.png"),
            ("FR-D2", "Driver posts GPS and status", "POST location; Collect/Deliver workflow"),
            ("FR-O1", "Office KPI overview", "GET /dashboard/overview"),
            ("FR-O2", "Office map with filters", "GET /dashboard/map"),
            ("FR-O3", "Office assigns driver", "POST assign"),
            ("FR-S1", "GPS/status simulation for demos", "POST /simulation/*"),
            ("FR-B1", "Shared API for all apps", "API_Contract_v1.docx; swagger.png"),
        ],
    )

    doc.add_heading("3. Non-functional requirements", level=1)
    add_table(
        doc,
        ("ID", "Requirement", "Result (local demo)"),
        [
            ("NFR-P1", "P95 API latency for overview under 300 ms", "4.69 ms (N=100, results.md)"),
            ("NFR-U1", "Heuristic review without blocking defects", "No major/critical (heuristic_review.md)"),
            ("NFR-S1", "Prototype bearer auth", "cmp600-demo-token header"),
            ("NFR-E1", "Ethics: simulated data only", "HE35; no field participants"),
            ("NFR-R1", "British English UI copy", "All three front ends"),
        ],
    )

    doc.add_heading("4. Out of scope (MVP)", level=1)
    for item in [
        "Real card payments and live customer accounts",
        "Production OAuth2, rate limiting, multi-tenant hosting",
        "OSRM or commercial road routing APIs",
        "WebSockets and Postgres in the submission build",
        "External user studies (SUS with participants)",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("5. Traceability", level=1)
    doc.add_paragraph(
        "Architecture Document.docx and API_Contract_v1.docx describe design. "
        "Evaluation_Plan_v1.docx describes tests. GitHub repository holds source and pytest logs."
    )

    out = DOC_DIR / "Requirements Document.docx"
    doc.save(out)
    print("Wrote", out.name)


def build_evaluation_plan() -> None:
    doc = Document()
    _style(doc)
    doc.add_heading("Evaluation Plan (Version 1)", 0)
    doc.add_paragraph(
        "This note describes how I tested the Door2Door CMP600 build before submission. "
        "Everything ran on my Windows laptop with the seeded SQLite database; I did not run a formal study "
        "with external users. The work was meant to support the dissertation evidence chapter and a short viva demo, "
        "not to sign off a production service."
    )

    doc.add_heading("1. What I needed to show", level=1)
    doc.add_paragraph(
        "First, the three apps had to complete the main logistics path: a client books a shipment, office staff "
        "see it on the map and assign a driver, the driver collects and delivers, and the client can still track status. "
        "Second, I had to compare API timing against NFR-P1 in the requirements register (P95 below 300 ms on the overview endpoint). "
        "Third, I walked through Nielsen heuristics on each front end so I could mention usability limits honestly in Chapter 6. "
        "I also kept screenshots, pytest output, and the latency log under Viva/ and Tests/ for Turnitin appendices."
    )
    doc.add_paragraph(
        "I deliberately did not load-test at cloud scale, run courier field trials, or calculate ROI. Those items sit outside the MVP brief."
    )

    doc.add_heading("2. Performance timing", level=1)
    doc.add_paragraph(
        "On 20 May 2026 I used Tests/performance/latency_check.py against GET /dashboard/overview while the API listened on "
        "127.0.0.1:8000. The script sent N=100 requests with the prototype bearer token cmp600-demo-token. "
        "Recorded values in results.md were mean 4.03 ms, P95 4.69 ms, and max 5.58 ms, so the run cleared the 300 ms bar with "
        "a wide margin. I treated the X-Response-Time-Ms response header as a sanity check against the client-side timings."
    )

    doc.add_heading("3. Usability pass", level=1)
    doc.add_paragraph(
        "On 21 May 2026 I spent roughly thirty-five minutes per app on a Nielsen-style heuristic review (Dashboard 5173, "
        "Client 5174, Driver 5175). Findings went into Tests/usability/heuristic_review.md with severities 0 to 4. "
        "I repeated the same four tasks on each build: sign-in, locate a delayed job on the office board, open tracking for shipment S2013, "
        "and read the driver stop order. No major or critical issues blocked those tasks; remaining notes were minor layout or list-length friction."
    )

    doc.add_heading("4. Cross-app checks", level=1)
    add_table(
        doc,
        ("Check", "What I looked for"),
        [
            ("Shared API", "cmp600-demo-token accepted on client, driver, and office builds"),
            ("Status line", "After driver marked Picked Up, client polling and office map showed the same label"),
            ("Simulation", "Starting GPS simulation moved the office map marker without manual SQL edits"),
            ("Build", "pytest reported 25 passed; each Vite app completed npm run build"),
        ],
    )

    doc.add_heading("5. Where the artefacts live", level=1)
    add_table(
        doc,
        ("Artefact", "Folder or file"),
        [
            ("Latency numbers", "Tests/performance/results.md"),
            ("Heuristic table", "Tests/usability/heuristic_review.md"),
            ("UI captures", "Viva/screenshots/"),
            ("Demo order", "Viva/DEMO_SCRIPT.md"),
        ],
    )

    doc.add_heading("6. Limits I state in the dissertation", level=1)
    doc.add_paragraph(
        "I am both developer and evaluator, which removes independence. Tests used localhost, SQLite, and ten-second polling rather than push events. "
        "Driver routing order is computed in the browser with haversine distance, not a commercial road engine. "
        "These results back the prototype claims in the report; they are not a warranty for a live courier network."
    )

    out = DOC_DIR / "Evaluation_Plan_v1.docx"
    doc.save(out)
    print("Wrote", out.name)


def build_cloud_deploy() -> None:
    doc = Document()
    _style(doc)
    doc.add_heading("Optional cloud hosting (Railway or Render)", 0)
    doc.add_paragraph(
        "CMP600 marking and my viva demo rely on the local stack started from server-control.bat in the repository root. "
        "The steps below are a personal deployment memo only, in case I publish the portfolio online after the module ends. "
        "Assessors can ignore this file unless they want to see how the Dockerised API might leave the laptop."
    )

    doc.add_heading("1. How I would split hosting", level=1)
    add_table(
        doc,
        ("Piece", "Typical host", "Notes"),
        [
            ("FastAPI backend", "Railway or Render web service", "Dockerfile under Source_Code/backend"),
            ("Three Vite front ends", "Vercel, Netlify, or static bucket", "Separate build per app"),
        ],
    )

    doc.add_heading("2. Railway backend", level=1)
    doc.add_paragraph(
        "I would link the GitHub repo, set the service root to Source_Code/backend, and let Railway build the supplied Dockerfile. "
        "The platform sets PORT; the container should start uvicorn with app.main:app bound to 0.0.0.0 on that port. "
        "API_BEARER_TOKEN can mirror the local demo token if I need the same auth story in the cloud."
    )
    doc.add_paragraph(
        "SQLite on Railway lives on ephemeral disk, so a redeploy wipes data unless I later swap to Postgres. "
        "Any public URL must still expose routes under /api/v1, for example https://your-service.up.railway.app/api/v1."
    )

    doc.add_heading("3. Front-end environment variables", level=1)
    doc.add_paragraph(
        "Each React app (dashboard, client_app, driver_app) needs VITE_API_URL pointing at that base before npm ci and npm run build. "
        "I would upload the three dist folders to static hosting. CORS is currently wide open in the prototype, which is acceptable for a demo but not for production."
    )

    doc.add_heading("4. Render variant", level=1)
    doc.add_paragraph(
        "Render can run the same backend folder with pip install -r requirements.txt and the same uvicorn start command. "
        "Front ends still read VITE_API_URL at build time; nothing else changes in the React code."
    )

    doc.add_heading("5. Quick checks before relying on cloud", level=1)
    doc.add_paragraph(
        "If I switch the demo to a remote API, I would hit GET /health, open /docs, and log into the client app against the cloud base URL. "
        "Ethics remain unchanged: seeded London data only, no real customer PII."
    )

    out = DOC_DIR / "Cloud_Deploy_Railway_Render.docx"
    doc.save(out)
    print("Wrote", out.name)


def md_to_docx(md_path: Path, docx_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    doc = Document()
    _style(doc)
    in_code = False
    code_lines: list[str] = []

    def flush_code() -> None:
        nonlocal code_lines
        if code_lines:
            doc.add_paragraph("\n".join(code_lines))
            code_lines = []

    for line in text.splitlines():
        if line.strip().startswith("```"):
            if in_code:
                flush_code()
            in_code = not in_code
            continue
        if in_code:
            code_lines.append(line)
            continue

        if not line.strip():
            flush_code()
            continue

        if line.startswith("# "):
            flush_code()
            doc.add_heading(line[2:].strip(), 0)
        elif line.startswith("## "):
            flush_code()
            doc.add_heading(line[3:].strip(), 1)
        elif line.startswith("### "):
            flush_code()
            doc.add_heading(line[4:].strip(), 2)
        elif line.startswith("|") and "|" in line[1:]:
            flush_code()
            doc.add_paragraph(line.strip())
        elif line.startswith("- ") or line.startswith("* "):
            flush_code()
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+\.\s", line):
            flush_code()
            doc.add_paragraph(re.sub(r"^\d+\.\s", "", line).strip(), style="List Number")
        else:
            flush_code()
            doc.add_paragraph(line.strip())

    flush_code()
    doc.save(docx_path)
    print("Wrote", docx_path.name)


def txt_to_docx(txt_path: Path, docx_path: Path) -> None:
    lines = txt_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    _style(doc)
    title = txt_path.stem.replace("_", " ")
    doc.add_heading(title, 0)
    for line in lines:
        if not line.strip():
            continue
        if line.isupper() and len(line) < 60:
            doc.add_heading(line.title(), 1)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(line.strip())
    doc.save(docx_path)
    print("Wrote", docx_path.name)


def main() -> None:
    build_requirements()
    build_evaluation_plan()
    build_cloud_deploy()

    skip_md = {"readme.md", "cloud_deploy_railway_render.md"}
    for md in DOC_DIR.glob("*.md"):
        if md.name.lower() in skip_md:
            continue
        md_to_docx(md, md.with_suffix(".docx"))

    for txt in DOC_DIR.glob("*.txt"):
        txt_to_docx(txt, txt.with_suffix(".docx"))

    print("Done.")


if __name__ == "__main__":
    main()
